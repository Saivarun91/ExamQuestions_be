"""Parse views: test parse (first 5), parse and save all."""
import io
import base64
import uuid
from rest_framework.decorators import api_view, parser_classes
from rest_framework.parsers import MultiPartParser, FormParser
from rest_framework.response import Response
from rest_framework import status
from django.views.decorators.csrf import csrf_exempt
from django.http import JsonResponse
from common.middleware import authenticate, restrict

from .helpers import (
    get_gemini_key,
    get_valid_gemini_model,
    build_parse_prompt,
    extract_questions_from_response,
    get_response_text,
)
from .models import ParsedInputQuestion


def _run_gemini_parse(settings_obj, full_prompt, file_content, file_ext, limit=None):
    """Call Gemini with prompt and file; return (questions_list, error_message)."""
    try:
        import google.generativeai as genai
    except ImportError:
        return [], "Gemini API not available. Install google-generativeai."

    key = get_gemini_key(settings_obj)
    if not key:
        return [], "GEMINI_API_KEY not found. Set it in AdminSettings or .env."
    genai.configure(api_key=key)
    model_name = get_valid_gemini_model(
        getattr(settings_obj, 'gemini_model_selector', None), genai
    )
    temp = float(getattr(settings_obj, 'temperature', 0))
    top_p = float(getattr(settings_obj, 'top_p', 1.0))

    try:
        model = genai.GenerativeModel(
            model_name if model_name.startswith("models/") else f"models/{model_name}"
        )
    except Exception:
        try:
            model = genai.GenerativeModel(model_name)
        except Exception as e:
            return [], f"Failed to load Gemini model: {e}"

    gen_config = {"temperature": temp, "top_p": top_p}
    try:
        if file_ext == "pdf":
            b64 = base64.b64encode(file_content).decode("utf-8")
            part = {"mime_type": "application/pdf", "data": b64}
            response = model.generate_content([full_prompt, part], generation_config=gen_config)
        else:
            response = model.generate_content(full_prompt, generation_config=gen_config)
    except Exception as e:
        return [], str(e)

    text = get_response_text(response)
    if not text or not text.strip():
        return [], "Empty response from Gemini."
    questions = extract_questions_from_response(text)
    if limit:
        questions = questions[: int(limit)]
    return questions, None


@api_view(['POST'])
@authenticate
@restrict(['admin'])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
def test_parse(request):
    """Parse first 5 questions only; return success and sample (no save)."""
    try:
        from settings_app.models import AdminSettings
        settings_obj = AdminSettings.objects.first() or AdminSettings()
        file = request.FILES.get("file")
        if not file:
            return JsonResponse({"success": False, "error": "No file provided"}, status=400)
        if not request.session.session_key:
            request.session.create()

        session_id = request.session.session_key
        instructions = request.POST.get("parsing_instructions", "")
        parsing_prompt = request.POST.get("parsing_prompt", "").strip()
        if not parsing_prompt:
            prompts = getattr(settings_obj, 'prompts', {}) or {}
            prompt1 = prompts.get('prompt1', {})
            parsing_prompt = (prompt1.get('prompt', '') if prompt1 else '') or ''
        if not parsing_prompt:
            return JsonResponse({"success": False, "error": "Please provide the parsing prompt."}, status=400)
        if not request.session.session_key:
            request.session.create()

        session_id = request.session.session_key
        file_content = file.read()
        file_ext = (file.name or "").split(".")[-1].lower()
        if file_ext == "docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_content))
                doc_text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                doc_text = file_content.decode("utf-8", errors="ignore")
        elif file_ext == "pdf":
            doc_text = None
        else:
            return JsonResponse({"success": False, "error": "Unsupported file type. Use PDF or DOCX."}, status=400)

        if file_ext == "pdf":
            full_prompt = build_parse_prompt(parsing_prompt, "[PDF attached]", instructions)
            questions, err = _run_gemini_parse(settings_obj, full_prompt, file_content, "pdf", limit=5)
        else:
            full_prompt = build_parse_prompt(parsing_prompt, doc_text, instructions)
            questions, err = _run_gemini_parse(settings_obj, full_prompt, doc_text, "docx", limit=5)

        if err:
            return JsonResponse({"success": False, "error": err}, status=500)
        return JsonResponse({
            "success": True,
            "parsed_count": len(questions),
            "questions": questions,
            "session_id": session_id, 
            "message": f"Parsing test successful: {len(questions)} question(s) parsed.",
        })
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=500)


@api_view(['POST'])
@authenticate
@restrict(['admin'])
@parser_classes([MultiPartParser, FormParser])
@csrf_exempt
def parse_save_all(request):
    """Parse full document with Gemini and save to Input Questions (question + options only)."""
    try:
        from settings_app.models import AdminSettings
        settings_obj = AdminSettings.objects.first() or AdminSettings()
        file = request.FILES.get("file")
        if not file:
            return JsonResponse({"success": False, "error": "No file provided"}, status=400)
        instructions = request.POST.get("parsing_instructions", "")
        parsing_prompt = request.POST.get("parsing_prompt", "").strip()
        if not parsing_prompt:
            prompts = getattr(settings_obj, 'prompts', {}) or {}
            prompt1 = prompts.get('prompt1', {})
            parsing_prompt = (prompt1.get('prompt', '') if prompt1 else '') or ''
        if not parsing_prompt:
            return JsonResponse({"success": False, "error": "Please provide the parsing prompt."}, status=400)

        file_content = file.read()
        file_ext = (file.name or "").split(".")[-1].lower()
        if file_ext == "docx":
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(io.BytesIO(file_content))
                doc_text = "\n".join(p.text for p in doc.paragraphs)
            except ImportError:
                doc_text = file_content.decode("utf-8", errors="ignore")
        elif file_ext == "pdf":
            doc_text = None
        else:
            return JsonResponse({"success": False, "error": "Unsupported file type. Use PDF or DOCX."}, status=400)

        if file_ext == "pdf":
            full_prompt = build_parse_prompt(parsing_prompt, "[PDF attached]", instructions)
            questions, err = _run_gemini_parse(settings_obj, full_prompt, file_content, "pdf")
        else:
            full_prompt = build_parse_prompt(parsing_prompt, doc_text, instructions)
            questions, err = _run_gemini_parse(settings_obj, full_prompt, doc_text, "docx")

        if err:
            return JsonResponse({"success": False, "error": err}, status=500)

        batch_id = str(uuid.uuid4())
        saved = 0
        if not request.session.session_key:
            request.session.create()

        session_id = request.session.session_key
        for q in questions:
            qt = (q.get("question_text") or "").strip()
            opts = q.get("options") or []
            if not isinstance(opts, list):
                opts = []
            opts = [o.get("text", str(o)) if isinstance(o, dict) else str(o) for o in opts]
            flag = (q.get("parsing_flag") or "VALID").strip().upper()
            if flag not in ("VALID", "INVALID"):
                flag = "VALID"
            if not qt:
                continue
            ParsedInputQuestion(question_text=qt, options=opts, parsing_flag=flag, batch_id=batch_id,session_id=session_id ).save()
            saved += 1
        return JsonResponse({
            "success": True,
            "parsed_count": len(questions),
            "saved_count": saved,
            "batch_id": batch_id,
            "session_id": session_id,
            "message": f"Parsed {len(questions)} question(s), saved {saved} to Input Questions.",
        })
    except Exception as e:
        return JsonResponse({"success": False, "error": str(e)}, status=500)
